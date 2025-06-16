import os, re
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, permissions
from rest_framework.decorators import api_view, permission_classes
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth import update_session_auth_hash
from .utils import calculate_statistics

from django.views.decorators.csrf import csrf_exempt

from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi
from .utils import calculate_statistics_for_collection
from .serializers import RegisterSerializer, LoginSerializer, ChangePasswordSerializer

from .authentication import CsrfExemptSessionAuthentication
from rest_framework.authentication import SessionAuthentication

from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator

from uuid import UUID
from rest_framework.permissions import IsAuthenticated
import uuid

from rest_framework.generics import CreateAPIView
from rest_framework import generics, permissions
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.exceptions import PermissionDenied
from .models import Document, Collection
from .serializers import DocumentSerializer, CollectionCreateSerializer
from .utils import calculate_tf_idf_for_document
from rest_framework.parsers import MultiPartParser, FormParser

from docx import Document as DocxDocument
import PyPDF2
from .utils import calculate_statistics
from .serializers import CollectionSerializer  # если у тебя уже есть сериализатор
from django.shortcuts import get_object_or_404
from .utils import calculate_tf_for_collection
from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi
from django.db.models import Count


# --- STATUS ---
@swagger_auto_schema(method='get', responses={200: 'Сервис работает'})
@api_view(['GET'])
def status_view(request):
    return Response({"status": "OK"})


# --- VERSION ---
@swagger_auto_schema(method='get', responses={200: 'Текущая версия'})
@api_view(['GET'])
def version_view(request):
    version = os.environ.get("APP_VERSION", "unknown")
    return Response({"version": version})


# --- METRICS ---
@swagger_auto_schema(method='get', responses={200: 'Метрики сервиса'})
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def metrics_view(request):
    import PyPDF2
    from docx import Document as DocxDocument

    user = request.user
    documents = Document.objects.filter(user=user)
    word_set = set()

    for doc in documents:
        try:
            file_path = doc.file.path
            ext = os.path.splitext(file_path)[1].lower()
            text = ""

            if ext == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif ext == ".pdf":
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = "".join(page.extract_text() or "" for page in reader.pages)
            elif ext == ".docx":
                docx = DocxDocument(file_path)
                text = "\n".join([p.text for p in docx.paragraphs])
            else:
                continue  # Пропустить неподдерживаемый формат

            words = re.findall(r'\b\w+\b', text.lower())
            word_set.update(words)

        except Exception as e:
            continue  # Пропустить при любой ошибке

    metrics = {
        "files_uploaded": documents.count(),
        "unique_words_analyzed": len(word_set)
    }

    return Response(metrics)


# --- REGISTER ---
class RegisterView(APIView):
    permission_classes = [permissions.AllowAny]

    @swagger_auto_schema(
        request_body=RegisterSerializer,
        responses={201: "Пользователь создан", 400: "Ошибка валидации"}
    )
    def post(self, request):
        serializer = RegisterSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Пользователь создан"}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# --- LOGIN ---
class LoginView(APIView):
    permission_classes = [permissions.AllowAny]

    @swagger_auto_schema(
        request_body=LoginSerializer,
        responses={200: "Вход выполнен", 400: "Ошибка валидации", 401: "Неверные данные"}
    )
    def post(self, request):
        serializer = LoginSerializer(data=request.data)
        if serializer.is_valid():
            user = authenticate(
                username=serializer.validated_data['username'],
                password=serializer.validated_data['password']
            )
            if user is not None:
                login(request, user)
                return Response({"message": "Вход выполнен", "user_id": user.id}, status=status.HTTP_200_OK)
            return Response({"error": "Неверные данные"}, status=status.HTTP_401_UNAUTHORIZED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# --- LOGOUT ---
class LogoutView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    @swagger_auto_schema(responses={200: "Выход выполнен"})
    def get(self, request):
        logout(request)
        return Response({"message": "Выход выполнен"}, status=status.HTTP_200_OK)


# --- CHANGE PASSWORD ---
class ChangePasswordView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    authentication_classes = [CsrfExemptSessionAuthentication]
    permission_classes = [permissions.IsAuthenticated]

    @swagger_auto_schema(
        manual_parameters=[
            openapi.Parameter('user_id', openapi.IN_PATH, description="UUID пользователя", type=openapi.TYPE_STRING,
                              format='uuid')
        ],
        request_body=ChangePasswordSerializer,
        responses={200: "Пароль изменён", 400: "Ошибка", 403: "Недостаточно прав"}
    )
    def patch(self, request, user_id: UUID):
        if str(request.user.id) != str(user_id):
            return Response({"error": "Недостаточно прав"}, status=status.HTTP_403_FORBIDDEN)

        serializer = ChangePasswordSerializer(data=request.data)
        if serializer.is_valid():
            if not request.user.check_password(serializer.validated_data['old_password']):
                return Response({"old_password": "Неверный текущий пароль"}, status=status.HTTP_400_BAD_REQUEST)
            request.user.set_password(serializer.validated_data['new_password'])
            request.user.save()
            update_session_auth_hash(request, request.user)
            return Response({"message": "Пароль изменён"}, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# --- DELETE USER ---
@method_decorator(csrf_exempt, name='dispatch')
class DeleteUserView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    @swagger_auto_schema(responses={200: "Пользователь удалён", 403: "Недостаточно прав"})
    def delete(self, request, user_id):
        if str(request.user.id) != str(user_id):
            return Response({"error": "Недостаточно прав"}, status=status.HTTP_403_FORBIDDEN)
        request.user.delete()
        logout(request)
        return Response({"message": "Пользователь удалён"}, status=status.HTTP_200_OK)


class DocumentListCreateView(generics.ListCreateAPIView):
    queryset = Document.objects.all()
    serializer_class = DocumentSerializer
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    @swagger_auto_schema(
        request_body=DocumentSerializer,
        responses={201: "Документ добавлен"},
        operation_description="Загрузка документа",
    )
    def post(self, request, *args, **kwargs):
        return self.create(request, *args, **kwargs)

    def get_queryset(self):
        return Document.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

class DocumentDetailView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, document_id):
        try:
            document = Document.objects.get(id=document_id, user=request.user)
            file_path = document.file.path

            ext = os.path.splitext(file_path)[1].lower()
            text = ""

            if ext == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif ext == ".pdf":
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = "".join(page.extract_text() for page in reader.pages if page.extract_text())
            elif ext == ".docx":
                doc = DocxDocument(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            else:
                return Response({"error": "Unsupported file format."}, status=status.HTTP_400_BAD_REQUEST)

            return Response({
                "id": str(document.id),
                "filename": os.path.basename(file_path),
                "content": text,
            })
        except Document.DoesNotExist:
            return Response({"error": "Document not found or access denied."}, status=status.HTTP_404_NOT_FOUND)

class DocumentStatisticsView(APIView):
    def get(self, request, document_id):
        document = get_object_or_404(Document, id=document_id)
        collections = document.collections.all()

        if not collections:
            return Response({"error": "Документ не входит ни в одну коллекцию"}, status=400)

        # Берём ВСЕ документы из коллекций, куда входит этот документ
        all_documents = Document.objects.filter(collections__in=collections).distinct()

        stats = calculate_statistics(document, all_documents)
        return Response(stats)



class DocumentDeleteView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def delete(self, request, document_id):
        try:
            document = Document.objects.get(id=document_id)
            if document.user != request.user:
                raise PermissionDenied("Нет доступа к этому документу")
        except Document.DoesNotExist:
            return Response({'error': 'Документ не найден'}, status=404)

        document.delete()
        return Response({'message': 'Документ удалён'})

class CollectionListView(APIView):
    def get(self, request):
        collections = Collection.objects.all()
        serializer = CollectionSerializer(collections, many=True)
        return Response(serializer.data)

class CollectionDetailView(APIView):
    def get(self, request, collection_id):
        collection = get_object_or_404(Collection, id=collection_id)
        document_ids = list(collection.documents.values_list('id', flat=True))
        return Response(document_ids)

class CollectionStatisticsView(APIView):
    def get(self, request, collection_id):
        collection = get_object_or_404(Collection, pk=collection_id)
        documents = collection.documents.all()

        if not documents.exists():
            return Response({"detail": "No documents found in the collection"}, status=404)

        stats = calculate_statistics_for_collection(documents)
        return Response(stats)



class CollectionAddDocumentView(APIView):
    def post(self, request, collection_id, document_id):
        collection = get_object_or_404(Collection, id=collection_id)
        document = get_object_or_404(Document, id=document_id)
        collection.documents.add(document)
        return Response({"status": "document added"})

class CollectionRemoveDocumentView(APIView):
    def delete(self, request, collection_id, document_id):
        collection = get_object_or_404(Collection, id=collection_id)
        document = get_object_or_404(Document, id=document_id)
        collection.documents.remove(document)
        return Response({"status": "document removed"})

class CollectionCreateView(CreateAPIView):
    queryset = Collection.objects.all()
    serializer_class = CollectionCreateSerializer

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

class DocumentHuffmanView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    @swagger_auto_schema(
        operation_description="Получить содержимое документа, закодированное кодом Хаффмана",
        responses={200: openapi.Response(description="Успешный ответ", examples={
            "application/json": {
                "id": "UUID документа",
                "filename": "название файла",
                "encoded": "хаффмановская строка"
            }
        }),
        404: "Документ не найден или нет доступа",
        400: "Неподдерживаемый формат файла",
        500: "Ошибка обработки"},
    )
    def get(self, request, document_id):
        try:
            document = Document.objects.get(id=document_id, user=request.user)
        except Document.DoesNotExist:
            return Response({'error': 'Документ не найден или доступ запрещён'}, status=404)

        file_path = document.file.path
        ext = os.path.splitext(file_path)[1].lower()
        text = ""

        try:
            if ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            elif ext == ".pdf":
                import PyPDF2
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = "".join(page.extract_text() or "" for page in reader.pages)
            elif ext == ".docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = "\\n".join([p.text for p in doc.paragraphs])
            else:
                return Response({"error": "Неподдерживаемый формат файла"}, status=400)

            from .utils import huffman_encode
            encoded = huffman_encode(text)

            return Response({
                "id": str(document.id),
                "filename": os.path.basename(file_path),
                "encoded": encoded
            })

        except Exception as e:
            return Response({"error": str(e)}, status=500)
